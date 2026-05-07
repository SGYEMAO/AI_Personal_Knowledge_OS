"""Content extractors for URLs and local files."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


class ExtractionError(Exception):
    """Raised when content extraction fails."""


ANTI_BOT_MESSAGE = "该网页可能存在反爬虫或访问验证，无法直接提取正文。请尝试复制正文、使用其他网页，或保存为 PDF 后再导入。"
ANTI_BOT_KEYWORDS = (
    "当前环境存在异常",
    "需要完成验证",
    "environment_exception",
    "access denied",
    "verify",
    "captcha",
)


class AntiBotDetectionError(ExtractionError):
    """Raised when a URL looks like an anti-bot or verification page."""


@dataclass(frozen=True)
class ExtractedContent:
    """Normalized extracted text and source metadata."""

    title: str
    text: str
    source_type: str
    source_path: str | None = None
    source_url: str | None = None


def extract_url(url: str) -> ExtractedContent:
    """Extract readable article text from a URL, preferring trafilatura."""
    if not url.lower().startswith(("http://", "https://")):
        raise ExtractionError("URL must start with http:// or https://")

    html = ""
    text = ""
    title = url

    try:
        import trafilatura

        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            html = downloaded
            extracted = trafilatura.extract(
                downloaded,
                include_comments=False,
                include_tables=True,
                favor_recall=True,
            )
            if extracted:
                text = extracted.strip()
            try:
                metadata = trafilatura.extract_metadata(downloaded)
                if metadata and metadata.title:
                    title = metadata.title.strip()
            except Exception:
                pass
    except Exception:
        text = ""

    if not text:
        title, text = _extract_url_with_bs4(url, html=html)

    if not text.strip():
        raise ExtractionError(f"No readable text found at URL: {url}")

    if looks_like_verification_page(text):
        raise AntiBotDetectionError(ANTI_BOT_MESSAGE)

    return ExtractedContent(title=title or url, text=text, source_type="url", source_url=url)


def looks_like_verification_page(text: str) -> bool:
    """Return True when extracted URL text looks like anti-bot verification content."""
    normalized = text.lower()
    return any(keyword.lower() in normalized for keyword in ANTI_BOT_KEYWORDS)


def _extract_url_with_bs4(url: str, html: str = "") -> tuple[str, str]:
    """Fallback URL extraction with requests and BeautifulSoup."""
    try:
        import requests
        from bs4 import BeautifulSoup

        if not html:
            response = requests.get(url, timeout=20, headers={"User-Agent": "AI-Knowledge-OS/0.1"})
            response.raise_for_status()
            html = response.text

        soup = BeautifulSoup(html, "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else url
        body = soup.get_text("\n", strip=True)
        return title, body
    except Exception as exc:
        raise ExtractionError(f"Failed to extract URL content: {exc}") from exc


def extract_file(path: Path) -> ExtractedContent:
    """Extract text from a supported local file."""
    source_path = str(path.expanduser().resolve())
    path = Path(source_path)
    if not path.exists() or not path.is_file():
        raise ExtractionError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = _read_text_file(path)
    elif suffix == ".pdf":
        text = _read_pdf_file(path)
    elif suffix == ".docx":
        text = _read_docx_file(path)
    else:
        raise ExtractionError(f"Unsupported file type: {suffix}")

    if not text.strip():
        raise ExtractionError(f"No text extracted from file: {path}")

    return ExtractedContent(
        title=path.stem,
        text=text,
        source_type=suffix.lstrip("."),
        source_path=source_path,
    )


def _read_text_file(path: Path) -> str:
    """Read a text or Markdown file with common encodings."""
    encodings = ("utf-8", "utf-8-sig", "cp1252")
    last_error: Exception | None = None
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        raise ExtractionError(f"Failed to read text file: {last_error or exc}") from exc


def _read_pdf_file(path: Path) -> str:
    """Read text from a PDF with PyMuPDF."""
    try:
        import fitz

        parts: list[str] = []
        with fitz.open(path) as document:
            for page in document:
                parts.append(page.get_text("text"))
        return "\n".join(parts)
    except Exception as exc:
        raise ExtractionError(f"Failed to read PDF file: {exc}") from exc


def _read_docx_file(path: Path) -> str:
    """Read paragraphs and tables from a DOCX file."""
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        raise ExtractionError(f"Failed to read DOCX file: {exc}") from exc
