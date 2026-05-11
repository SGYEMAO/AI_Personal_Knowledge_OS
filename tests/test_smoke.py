"""Smoke tests for the local MVP without calling the OpenAI API."""

from __future__ import annotations

import http.server
import os
import shutil
import socketserver
import threading
import unittest
import uuid
from functools import partial
from pathlib import Path

from src import ingest
from src.extractors import ANTI_BOT_MESSAGE
from src.storage import find_existing_document
from src.utils import load_settings


def fake_summary(text: str, **kwargs: object) -> dict[str, object]:
    """Return a deterministic structured summary for tests."""
    return {
        "title": kwargs.get("title_hint") or "Test Title",
        "summary": text[:120],
        "key_points": ["Point one", "Point two"],
        "topics": ["tests"],
        "action_items": [],
        "entities": ["Codex"],
    }


class KnowledgeOSSmokeTests(unittest.TestCase):
    """Verify extraction, persistence, and deduplication with local fixtures."""

    def setUp(self) -> None:
        tmp_parent = Path(__file__).resolve().parent / "_tmp"
        tmp_parent.mkdir(parents=True, exist_ok=True)
        self.root = tmp_parent / f"run_{uuid.uuid4().hex}"
        self.root.mkdir(parents=True, exist_ok=True)
        os.environ["OPENAI_API_KEY"] = "test-key"
        os.environ["KNOWLEDGE_BASE_PATH"] = str(self.root / "knowledge_base")
        os.environ["SQLITE_DB_PATH"] = str(self.root / "data" / "knowledge.db")
        os.environ["REPORTS_PATH"] = str(self.root / "reports")
        os.environ["EXPORTS_PATH"] = str(self.root / "exports")
        os.environ["VECTOR_STORE_PATH"] = str(self.root / "vector_store")
        os.environ["LLM_PROVIDER"] = "openai"
        os.environ["OLLAMA_BASE_URL"] = "http://localhost:11434"
        os.environ["OLLAMA_MODEL"] = "llama3.1"
        os.environ["EMBEDDING_PROVIDER"] = "ollama"
        os.environ["OLLAMA_EMBED_MODEL"] = "nomic-embed-text"
        self.settings = load_settings(create_dirs=True)
        self.original_summarize = ingest.summarize_text
        self.original_upsert = ingest.upsert_document_embedding
        self.embedding_calls: list[dict[str, object]] = []
        ingest.summarize_text = fake_summary
        ingest.upsert_document_embedding = self._fake_upsert

    def tearDown(self) -> None:
        ingest.summarize_text = self.original_summarize
        ingest.upsert_document_embedding = self.original_upsert
        shutil.rmtree(self.root, ignore_errors=True)

    def test_process_txt_md_docx_pdf_and_dedupe(self) -> None:
        """Process supported local files and avoid duplicate inserts."""
        source_dir = self.root / "sources"
        source_dir.mkdir()
        (source_dir / "note.txt").write_text("Text note content for testing.", encoding="utf-8")
        (source_dir / "memo.md").write_text("# Memo\nMarkdown content for testing.", encoding="utf-8")
        self._write_docx(source_dir / "brief.docx")
        self._write_pdf(source_dir / "paper.pdf")

        results = ingest.process_folder(source_dir, settings=self.settings)
        self.assertEqual(len(results), 4)
        self.assertTrue(all(result.status == "processed" for result in results), results)
        self.assertTrue(all(result.markdown_path and result.markdown_path.exists() for result in results))
        self.assertTrue(self.settings.sqlite_db_path.exists())

        duplicate = ingest.process_file(source_dir / "note.txt", settings=self.settings)
        self.assertEqual(duplicate.status, "already_processed")
        self.assertIsNotNone(duplicate.markdown_path)

    def test_process_url(self) -> None:
        """Process a local HTTP URL using the URL extractor path."""
        web_dir = self.root / "web"
        web_dir.mkdir()
        (web_dir / "index.html").write_text(
            "<html><head><title>Local Article</title></head>"
            "<body><article><h1>Local Article</h1><p>URL content for testing.</p></article></body></html>",
            encoding="utf-8",
        )

        handler = partial(http.server.SimpleHTTPRequestHandler, directory=str(web_dir))
        with socketserver.TCPServer(("127.0.0.1", 0), handler) as server:
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            url = f"http://127.0.0.1:{server.server_address[1]}/index.html"
            result = ingest.process_url(url, settings=self.settings)
            server.shutdown()
            thread.join(timeout=5)

        self.assertEqual(result.status, "processed", result)
        existing = find_existing_document(self.settings.sqlite_db_path, source_url=url)
        self.assertIsNotNone(existing)

    def test_url_verification_page_does_not_persist_or_summarize(self) -> None:
        """Block anti-bot verification pages before LLM and persistence."""

        def fail_summary(text: str, **kwargs: object) -> dict[str, object]:
            raise AssertionError("LLM should not be called for verification pages")

        ingest.summarize_text = fail_summary
        web_dir = self.root / "blocked_web"
        web_dir.mkdir()
        (web_dir / "index.html").write_text(
            "<html><head><title>Verify</title></head>"
            "<body><main>environment_exception verification required captcha verify</main></body></html>",
            encoding="utf-8",
        )

        handler = partial(http.server.SimpleHTTPRequestHandler, directory=str(web_dir))
        with socketserver.TCPServer(("127.0.0.1", 0), handler) as server:
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            url = f"http://127.0.0.1:{server.server_address[1]}/index.html"
            result = ingest.process_url(url, settings=self.settings)
            server.shutdown()
            thread.join(timeout=5)

        self.assertEqual(result.status, "blocked", result)
        self.assertEqual(result.message, ANTI_BOT_MESSAGE)
        self.assertFalse(self.settings.sqlite_db_path.exists())
        self.assertEqual(list(self.settings.knowledge_base_path.rglob("*.md")), [])
        self.assertEqual(self.embedding_calls, [])

    def test_failed_extraction_does_not_persist_or_summarize(self) -> None:
        """Failed local extraction returns failed without LLM, SQLite, or Markdown."""

        def fail_summary(text: str, **kwargs: object) -> dict[str, object]:
            raise AssertionError("LLM should not be called for failed extraction")

        ingest.summarize_text = fail_summary
        source_dir = self.root / "empty_source"
        source_dir.mkdir()
        empty_file = source_dir / "empty.txt"
        empty_file.write_text("", encoding="utf-8")

        result = ingest.process_file(empty_file, settings=self.settings)

        self.assertEqual(result.status, "failed", result)
        self.assertFalse(self.settings.sqlite_db_path.exists())
        self.assertEqual(list(self.settings.knowledge_base_path.rglob("*.md")), [])
        self.assertEqual(self.embedding_calls, [])

    def test_processed_document_calls_semantic_upsert(self) -> None:
        """Processed documents are incrementally added to the semantic index."""
        source_file = self.root / "semantic.txt"
        source_file.write_text("Semantic indexing test content.", encoding="utf-8")

        result = ingest.process_file(source_file, settings=self.settings)

        self.assertEqual(result.status, "processed", result)
        self.assertEqual(result.semantic_index_status, "indexed")
        self.assertEqual(len(self.embedding_calls), 1)
        self.assertEqual(self.embedding_calls[0]["id"], 1)

    def test_embedding_failure_does_not_fail_ingest(self) -> None:
        """Embedding failures are surfaced without failing ingest."""

        def fail_upsert(doc: dict[str, object], settings: object) -> None:
            raise RuntimeError("embedding service down")

        ingest.upsert_document_embedding = fail_upsert
        source_file = self.root / "embedding_fail.txt"
        source_file.write_text("Content still saves even if embeddings fail.", encoding="utf-8")

        result = ingest.process_file(source_file, settings=self.settings)

        self.assertEqual(result.status, "processed", result)
        self.assertEqual(result.semantic_index_status, "failed")
        self.assertIn("embedding service down", result.semantic_index_error)
        self.assertTrue(result.markdown_path and result.markdown_path.exists())
        self.assertTrue(self.settings.sqlite_db_path.exists())

    def _write_docx(self, path: Path) -> None:
        """Create a small DOCX fixture if python-docx is installed."""
        try:
            from docx import Document
        except ImportError as exc:
            self.skipTest(f"python-docx not installed: {exc}")
        document = Document()
        document.add_heading("Brief", level=1)
        document.add_paragraph("DOCX content for testing.")
        document.save(path)

    def _write_pdf(self, path: Path) -> None:
        """Create a small PDF fixture if PyMuPDF is installed."""
        try:
            import fitz
        except ImportError as exc:
            self.skipTest(f"PyMuPDF not installed: {exc}")
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "PDF content for testing.")
        document.save(path)
        document.close()

    def _fake_upsert(self, doc: dict[str, object], settings: object) -> None:
        """Record semantic upsert calls without invoking Ollama or ChromaDB."""
        self.embedding_calls.append(doc)


if __name__ == "__main__":
    unittest.main()
